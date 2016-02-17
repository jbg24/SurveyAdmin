# Generate login codes for Qualtrics panel
from docx import Document
from docx.shared import Inches, Pt
import pandas as pd
import math
import subprocess
import shutil

def make_doc(title,login_code_list):
    '''
    '''

    document = Document()

    # Styling for everything but header
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(110)

    # Create heading
    paragraph = document.add_paragraph(title)
    paragraph.style = 'Heading3'
    paragraph.style.font.size = Pt(20)

    # Create table and split login code list
    table_length = int(math.ceil(len(login_code_list['LoginCodes'])/float(2)))
    login_code_list_part1 = login_code_list['LoginCodes'][:table_length]
    login_code_list_part2 = login_code_list['LoginCodes'][table_length:]

    # Fill in table
    table = document.add_table(rows=table_length, cols=2)
    for i in range(0,table_length):
        row = table.rows[i].cells
        row[0].text = login_code_list_part1[i]
        if i < len(login_code_list_part2):
            row[1].text = login_code_list_part2[i]

    #hdr_cells = table.rows[0].cells
    #hdr_cells[0].text = 'Qty'
    #hdr_cells[1].text = 'Id'
    #hdr_cells[2].text = 'Desc'

    #document.add_page_break()
    document.save('demo.doc')

def id_generator(size=4):
    '''
    Generate random combination of 1 letter and 3 digits
    '''

    #remove ambiguous o and l characters
    replace_chars = ["o", "l"]
    chars = string.ascii_lowercase

    for char in replace_chars:
        chars = chars.replace(char, "")

    #remove 0
    nums =  string.digits.replace("0","")

    #create code: e.g. b452
    rand_nums = ''.join(random.choice(nums) for _ in range(size-1))
    return random.choice(chars) + rand_nums

def get_login_codes(enrollment):
    '''
    Return list of randomly generate login codes in format
    :param num:
    :return:
    '''

    # Checks to make sure won't enter infinite loop, default combinations is 17496 so check at 15000
    if enrollment > 15000:
        print "WARNING: Large number of students (over 15000). Will now terminate"
        sys.exit()

    # Random and unique login codes - use a set
    random_list = set()

    # Create the same number of login codes as the number of enrolled or what we set above as enrollment
    while (len(random_list) < enrollment):
        random_list.add(id_generator())

    return list(random_list)

def copy2(wb):
    '''
    Patch: add this function to the end of xlutils/copy.py
    :param wb: excel workbook using xlrd
    :return: workbook with format preserved
    '''
    w = XLWTWriter()
    process(
        XLRDReader(wb,'unknown.xls'),
        w
        )
    return w.output[0][1], w.style_list

def print_ose_login_codes(data_points,codes_path):
    '''
    Print the login codes for OSE - right now, only the ExternalDataReference field in the panel --> need to make a PDF
    Print two files - one with
    :param data_points: (Series)
    :param codes_path: (String)
    :return:
    '''
    length_df = len(data_points)
    main_file_length = int(length_df*0.6)
    data_points[:main_file_length].to_csv(codes_path,index=False,header=True)

    main_file_path = os.path.split(codes_path)[0]
    backup_file_name = os.path.split(codes_path)[1].split('.csv')[0] + '_backup.csv'
    backup_path = os.path.join(main_file_path,backup_file_name)
    data_points[main_file_length:].to_csv(backup_path,index=False,header=True)

def print_fft_login_codes(school_name, data_points,codes_path):
    '''
    Using login codes template, print login codes
    :param data_points:
    :param codes_path:
    :return:
    '''

    wbk = xlrd.open_workbook("TemplatesMapsImages/LoginCodeTemplate.xls", formatting_info=True)
    rdsheet = wbk.sheet_by_index(0)
    wbk_copy,style_list = copy2(wbk)

    sheet_cover = wbk_copy.get_sheet(0)
    xf_index = rdsheet.cell_xf_index(3, 3)
    sheet_cover.write(9,0,school_name + " - Login Codes",style_list[xf_index])

    row_num = 11
    for index, row in data_points.iterrows():
        xf_index = rdsheet.cell_xf_index(2, 2)
        sheet_cover.write(row_num,2,row['StudentID'],style_list[xf_index])
        sheet_cover.write(row_num,4,row['ExternalDataReference'],style_list[xf_index])
        row_num += 1


    sheet_cover.insert_bitmap('TemplatesMapsImages/YTimage.bmp',0,2)
    output_path = os.path.join(codes_path,school_name + '_logincodes.xls')
    wbk_copy.save(output_path)

def doc_to_pdf(doc_path):
    '''
    Save word document as PDF
    doc_path: (string) path to word doc_path
    '''
    input_filename = doc_path
    output_filename = 'output.pdf'

    p = subprocess.Popen(['unoconv', '--stdout', input_filename], stdout=subprocess.PIPE)
    with open(output_filename, 'w') as output:
       shutil.copyfileobj(p.stdout, output)



if __name__ == '__main__':
    testing_header = 'Test Header'
    test_list = pd.Series({'LoginCodes':['abc','efs2','sdf','234','sdf','def','efg','hij','jkl']})
    make_doc(testing_header,test_list)
    #doc_to_pdf('demo.doc')
